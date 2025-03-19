"""
JSON to DOCX Converter
Copyright (C) 2024

This program is free software: you can redistribute it and/or modify
it under the terms of the GNU General Public License as published by
the Free Software Foundation, either version 3 of the License, or
(at your option) any later version.

This program is distributed in the hope that it will be useful,
but WITHOUT ANY WARRANTY; without even the implied warranty of
MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE. See the
GNU General Public License for more details.

You should have received a copy of the GNU General Public License
along with this program. If not, see <https://www.gnu.org/licenses/>.
"""

from fastapi import FastAPI, Request, HTTPException
from fastapi.responses import FileResponse
from docx import Document
import json
import tempfile
import os
from fastapi.middleware.cors import CORSMiddleware

app = FastAPI()

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.post("/convert")
async def convert_json_to_docx(request: Request):
    try:
        # Get the raw data and parse it
        raw_data = await request.body()
        data = json.loads(raw_data)
        
        doc = Document()
        doc.add_heading('JSON to DOCX Conversion', 0)
        
        # Handle nested dictionaries and lists
        def add_content(content, level=0):
            indent = "    " * level
            if isinstance(content, dict):
                for key, value in content.items():
                    if isinstance(value, (dict, list)):
                        doc.add_paragraph(f"{indent}{key}:")
                        add_content(value, level + 1)
                    else:
                        doc.add_paragraph(f"{indent}{key}: {value}")
            elif isinstance(content, list):
                for item in content:
                    add_content(item, level)
            else:
                doc.add_paragraph(f"{indent}{content}")
        
        add_content(data)
        
        # Create temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        
        return FileResponse(
            temp_file.name,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename="converted.docx"
        )
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid JSON format")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/")
async def serve_index():
    return FileResponse("index.html")